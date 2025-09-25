import streamlit as st
import traceback
from config.database import save_error
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from io import BytesIO

class ResumeBuilder:
    def __init__(self):
        """
        Initialize templates
        """
        self.templates = {
            "Classic": self.build_resume_template_1,
            "Modern Split": self.build_resume_template_2,
            "Block Style": self.build_resume_template_3
        }

    def generate_resume(self, data):
        """
        Generate a resume
        """
        try:

            # Create a new document
            doc = Document()

            # Get selected template name
            template_name = data.get("selected_template", "Classic")
            
            # Pick the right builder
            template_func = self.templates.get(template_name)

            if template_func:
                doc = template_func(doc, data)
            else:
                st.warning(f"⚠️ Unknown template \"{template_name}\", falling back to Classic.")
                doc = self.templates["Classic"](doc, data)
                        
            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except Exception:
            error = traceback.format_exc()
            save_error(error)

    def format_list_items(self, items):
        """
        Helper function to handle both string and list inputs
        """
        if isinstance(items, str):
            return [item.strip() for item in items.split("\n") if item.strip()]
        elif isinstance(items, list):
            return [item.strip() for item in items if item and item.strip()]
        return []
    
    def apply_styles(self, doc):
        """
        Define and register custom styles for the document
        """
        styles = doc.styles

        # Header style
        header_style = (styles.add_style("Header", WD_STYLE_TYPE.PARAGRAPH)
                            if "Header" not in styles 
                            else styles["Header"]
        )
        header_style.font.size = Pt(28)
        header_style.font.bold = True
        header_style.paragraph_format.space_after = Pt(4)
        header_style.font.name = "Calibri"

        # Section style
        section_style = (styles.add_style("Section", WD_STYLE_TYPE.PARAGRAPH)
                            if "Section" not in styles 
                            else styles["Section"]
        )
        section_style.font.size = Pt(18)
        section_style.font.bold = True
        section_style.font.color.rgb = RGBColor(0, 120, 215)
        section_style.paragraph_format.space_before = Pt(12)
        section_style.paragraph_format.space_after = Pt(6)
        section_style.font.name = "Calibri"

        # Normal text style
        normal_style = (styles.add_style("Normal Text", WD_STYLE_TYPE.PARAGRAPH)
                            if "Normal Text" not in styles 
                            else styles["Normal Text"]
        )
        normal_style.font.size = Pt(12)
        normal_style.font.name = "Calibri"
        normal_style.paragraph_format.space_after = Pt(2)

        # Contact style
        contact_style = (styles.add_style("Contact", WD_STYLE_TYPE.PARAGRAPH) 
                            if "Contact" not in styles 
                            else styles["Contact"]
        )
        contact_style.font.size = Pt(10)
        contact_style.font.name = "Calibri"
        contact_style.paragraph_format.space_after = Pt(6)

    def build_resume_template_1(self, doc, data):
        """
        Build resume using template 1 (Classic)
        """
        try:
            self.apply_styles(doc)

            # Add full name at the top
            name = doc.add_paragraph(data["personal_info"]["name"], style="Header")
            name.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Add contact in a single line
            contact_parts = []
            if data["personal_info"].get("email"): 
                contact_parts.append(data["personal_info"]["email"])
            if data["personal_info"].get("phone"): 
                contact_parts.append(data["personal_info"]["phone"])

            if contact_parts:
                doc.add_paragraph(" | ".join(contact_parts), style="Contact")

            # Add links in a single line
            link_parts = []
            if data["personal_info"].get("github"): 
                link_parts.append(data["personal_info"]["github"])
            if data["personal_info"].get("linkedin"): 
                link_parts.append(data["personal_info"]["linkedin"])
            if data["personal_info"].get("portfolio"): 
                link_parts.append(data["personal_info"]["portfolio"])            
            
            if link_parts:
                doc.add_paragraph(" | ".join(link_parts), style="Contact")

            # Summary
            if data.get("summary"):
                doc.add_paragraph("SUMMARY", style="Section")
                doc.add_paragraph(data["summary"], style="Normal Text")

            # Education
            if data.get("education"):
                doc.add_paragraph("EDUCATION", style="Section")
                for edu in data["education"]:
                    p = doc.add_paragraph(style="Normal Text")
                    p.add_run(f"{edu['school']}").bold = True
                    p.add_run(f"\n{edu['degree']} in {edu['field']}")
                    p.add_run(f" | Graduation: {edu['graduation_date']}")
                    
                    if edu.get("gpa"):
                        p.add_run(f" | GPA: {edu['gpa']}")

                    if edu.get("achievements"):
                        for achv in self.format_list_items(edu["achievements"]):
                            bullet = doc.add_paragraph(f"• {achv}", style="Normal Text")
                            bullet.paragraph_format.left_indent = Inches(0.3)

            # Experience
            if data.get("experience"):
                doc.add_paragraph("EXPERIENCE", style="Section")
                for exp in data["experience"]:
                    p = doc.add_paragraph(style="Normal Text")
                    p.add_run(f"{exp['position']} at {exp['company']}").bold = True
                    p.add_run(f" | {exp['start_date']} - {exp['end_date']}")

                    if exp.get("description"):
                        doc.add_paragraph(exp["description"], style="Normal Text")

                    if exp.get("responsibilities"):
                        for resp in self.format_list_items(exp["responsibilities"]):
                            bullet = doc.add_paragraph(f"• {resp}", style="Normal Text")
                            bullet.paragraph_format.left_indent = Inches(0.3)

            # Projects
            if data.get("projects"):
                doc.add_paragraph("PROJECTS", style="Section")
                for proj in data["projects"]:
                    p = doc.add_paragraph(style="Normal Text")
                    p.add_run(proj["name"]).bold = True
                    
                    if proj.get("technologies"):
                        p.add_run(f" | {proj['technologies']}")

                    if proj.get("link"):
                        doc.add_paragraph(f"Link: {proj['link']}", style="Contact")
                    
                    if proj.get("description"):
                        doc.add_paragraph(proj["description"], style="Normal Text")
                    
                    if proj.get("responsibilities"):
                        for resp in self.format_list_items(proj["responsibilities"]):
                            bullet = doc.add_paragraph(f"• {resp}", style="Normal Text")
                            bullet.paragraph_format.left_indent = Inches(0.3)

            # References
            if data.get("references"):
                doc.add_paragraph("REFERENCES", style="Section")
                for ref in data["references"]:

                    parts = [
                        ref.get("name", ""),
                        ref.get("company", ""),
                        ref.get("phone", ""),
                        ref.get("email", "")
                    ]

                    # Add all in a single line
                    p = doc.add_paragraph(style="Contact")
                    p.add_run(" | ".join([par for par in parts if par]))

            # Skills
            if data.get("skills"):
                doc.add_paragraph("SKILLS", style="Section")
                skills = data["skills"]
                
                def add_skill_category(category_name, title):
                    if skills.get(category_name):
                        p = doc.add_paragraph(style="Normal Text")
                        p.add_run(f"{title}: ").bold = True
                        p.add_run(", ".join(self.format_list_items(skills[category_name])))
                
                add_skill_category("technical", "Technical Skills")
                add_skill_category("soft", "Soft Skills")
                add_skill_category("languages", "Languages")
                add_skill_category("tools", "Tools & Technologies")

            # Margins
            for section in doc.sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.7)
                section.right_margin = Inches(0.7)

            return doc

        except Exception:
            error = traceback.format_exc()
            save_error(error)

    def build_resume_template_2(self, doc, data):
        """
        Build resume using template 2 (Modern Split)
        """
        try:
            self.apply_styles(doc)

            # Add full name at the top
            name = doc.add_paragraph(data["personal_info"]["name"], style="Header")
            name.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Create a table with 2 columns
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False

            left_col_width = Inches(2.5)
            right_col_width = Inches(4.0)
            table.columns[0].width = left_col_width
            table.columns[1].width = right_col_width

            left_cell = table.cell(0, 0)
            right_cell = table.cell(0, 1)

            # Remove borders
            for row in table.rows:
                for cell in row.cells:
                    for side in cell._element.xpath(".//w:tcBorders/*"):
                        side.getparent().remove(side)

            # -----------
            # Left column
            # -----------

            # Add contacts and links one after the other
            contact_parts = []
            if data["personal_info"].get("email"): 
                contact_parts.append(data["personal_info"]["email"])
            if data["personal_info"].get("phone"): 
                contact_parts.append(data["personal_info"]["phone"])
            if data["personal_info"].get("github"): 
                contact_parts.append(data["personal_info"]["github"])
            if data["personal_info"].get("linkedin"): 
                contact_parts.append(data["personal_info"]["linkedin"])
            if data["personal_info"].get("portfolio"): 
                contact_parts.append(data["personal_info"]["portfolio"])
            
            if contact_parts:
                for contact in contact_parts:
                    left_cell.add_paragraph(contact, style="Contact")

            # Summary
            if data.get("summary"):
                left_cell.add_paragraph("SUMMARY", style="Section")
                left_cell.add_paragraph(data["summary"], style="Normal Text")

            # References
            if data.get("references"):
                left_cell.add_paragraph("REFERENCES", style="Section")
                for ref in data["references"]:

                    parts = [
                        ref.get("name", ""),
                        ref.get("company", ""),
                        ref.get("phone", ""),
                        ref.get("email", "")
                    ]

                    # Add all in a single line
                    p = left_cell.add_paragraph(style="Contact")
                    p.add_run(" | ".join([par for par in parts if par]))

            # ------------
            # Right column
            # ------------

            # Education
            if data.get("education"):
                right_cell.add_paragraph("EDUCATION", style="Section")
                for edu in data["education"]:
                    p = right_cell.add_paragraph(style="Normal Text")
                    p.add_run(f"{edu['school']}").bold = True
                    p.add_run(f"\n{edu['degree']} in {edu['field']}")
                    p.add_run(f" | Graduation: {edu['graduation_date']}")

                    if edu.get("gpa"):
                        p.add_run(f" | GPA: {edu['gpa']}")

                    if edu.get("achievements"):
                            for achv in self.format_list_items(edu["achievements"]):
                                bullet = right_cell.add_paragraph(f"• {achv}", style="Normal Text")
                                bullet.paragraph_format.left_indent = Inches(0.3)

            # Experience
            if data.get("experience"):
                right_cell.add_paragraph("EXPERIENCE", style="Section")
                for exp in data["experience"]:
                    p = right_cell.add_paragraph(style="Normal Text")
                    p.add_run(f"{exp['position']} at {exp['company']}").bold = True
                    p.add_run(f" | {exp['start_date']} - {exp['end_date']}")
                    
                    if exp.get("description"):
                        right_cell.add_paragraph(exp["description"], style="Normal Text")
                    
                    if exp.get("responsibilities"):
                        for resp in self.format_list_items(exp["responsibilities"]):
                            bullet = right_cell.add_paragraph(f"• {resp}", style="Normal Text")
                            bullet.paragraph_format.left_indent = Inches(0.3)

            # Projects
            if data.get("projects"):
                right_cell.add_paragraph("PROJECTS", style="Section")
                for proj in data["projects"]:
                    p = right_cell.add_paragraph(style="Normal Text")
                    p.add_run(proj["name"]).bold = True

                    if proj.get("technologies"):
                        p.add_run(f" | {proj['technologies']}")
                    
                    if proj.get("link"):
                        right_cell.add_paragraph(f"Link: {proj['link']}", style="Contact")

                    if proj.get("description"):
                        right_cell.add_paragraph(proj["description"], style="Normal Text")
                    
                    if proj.get("responsibilities"):
                        for resp in self.format_list_items(proj["responsibilities"]):
                            bullet = right_cell.add_paragraph(f"• {resp}", style="Normal Text")
                            bullet.paragraph_format.left_indent = Inches(0.3)

            # Skills
            if data.get("skills"):
                right_cell.add_paragraph("SKILLS", style="Section")
                skills = data["skills"]

                def add_skill_category(category_name, title):
                    if skills.get(category_name):
                        p = right_cell.add_paragraph(style="Normal Text")
                        p.add_run(f"{title}: ").bold = True
                        p.add_run(", ".join(self.format_list_items(skills[category_name])))
                
                add_skill_category("technical", "Technical Skills")
                add_skill_category("soft", "Soft Skills")
                add_skill_category("languages", "Languages")
                add_skill_category("tools", "Tools & Technologies")

            # Margins
            for section in doc.sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

            return doc

        except Exception:
            error = traceback.format_exc()
            save_error(error)

    def build_resume_template_3(self, doc, data):
        """
        Build resume using template 3 (Block Style)
        """
        try:
            self.apply_styles(doc)

            # Add full name at the top
            name = doc.add_paragraph(data["personal_info"]["name"], style="Header")
            name.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Add contact in a single line
            contact_parts = []
            if data["personal_info"].get("email"): 
                contact_parts.append(data["personal_info"]["email"])
            if data["personal_info"].get("phone"): 
                contact_parts.append(data["personal_info"]["phone"])

            if contact_parts:
                doc.add_paragraph(" | ".join(contact_parts), style="Contact")
            
            # Add links in a single line
            link_parts = []
            if data["personal_info"].get("github"): 
                link_parts.append(data["personal_info"]["github"])
            if data["personal_info"].get("linkedin"): 
                link_parts.append(data["personal_info"]["linkedin"])
            if data["personal_info"].get("portfolio"): 
                link_parts.append(data["personal_info"]["portfolio"])            
            
            if link_parts:
                doc.add_paragraph(" | ".join(link_parts), style="Contact")

            # Helper function to create shaded block
            def add_block(title, content_list):

                # Create a table with 1 cell to simulate a block
                table = doc.add_table(rows=1, cols=1)
                table.autofit = False
                cell = table.cell(0, 0)
                cell.width = Inches(6.5)
                
                # Add shading
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), "D9EAF7")
                tcPr.append(shd)

                # Section title
                cell_paragraph = cell.add_paragraph(title, style="Section")
                cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Content
                for content in content_list:
                    p = cell.add_paragraph(content, style="Normal Text")
                    p.paragraph_format.space_after = Pt(2)

                # space between blocks
                doc.add_paragraph()

            # Summary
            if data.get("summary"):
                add_block("SUMMARY", [data["summary"]])

            # Education
            if data.get("education"):
                edu_list = []
                for edu in data["education"]:
                    line = f"{edu['school']} - {edu['degree']} in {edu['field']} | Graduation: {edu['graduation_date']}"
                    
                    if edu.get("gpa"):
                        line += f" | GPA: {edu['gpa']}"
                    edu_list.append(line)

                    if edu.get("achievements"):
                        for achv in self.format_list_items(edu["achievements"]):
                            edu_list.append(f"  • {achv}")
                    
                add_block("EDUCATION", edu_list)

            # Experience
            if data.get("experience"):
                exp_list = []
                for exp in data["experience"]:
                    exp_list.append(f"{exp['position']} at {exp['company']} | {exp['start_date']} - {exp['end_date']}")
                    
                    if exp.get("description"):
                        exp_list.append(f"  {exp['description']}")

                    if exp.get("responsibilities"):
                        for resp in self.format_list_items(exp["responsibilities"]):
                            exp_list.append(f"  • {resp}")

                add_block("EXPERIENCE", exp_list)

            # Projects
            if data.get("projects"):
                proj_list = []
                for proj in data["projects"]:
                    line = proj["name"]
                    
                    if proj.get("technologies"):
                        line += f" | {proj['technologies']}"
                    proj_list.append(line)
                    
                    if proj.get("link"):
                        proj_list.append(f"Link: {proj['link']}")
                    
                    if proj.get("description"):
                        proj_list.append(f"  {proj['description']}")
                    
                    if proj.get("responsibilities"):
                        for resp in self.format_list_items(proj["responsibilities"]):
                            proj_list.append(f"  • {resp}")

                add_block("PROJECTS", proj_list)

            # References
            if data.get("references"):
                ref_list = []
                for ref in data["references"]:
                    line = ref["name"]

                    if ref.get("company"):
                        line += f" | {ref['company']}"
                    ref_list.append(line)

                    if ref.get("phone"):
                        ref_list.append(f"  {ref['phone']}")

                    if ref.get("email"):
                        ref_list.append(f"  {ref['email']}")

                add_block("REFERENCES", ref_list)

            # Skills
            if data.get("skills"):
                skills = data["skills"]
                skills_list = []
                
                def add_skill_category(category_name, title):
                    if skills.get(category_name):
                        skills_list.append(f"{title}: {', '.join(self.format_list_items(skills[category_name]))}")

                add_skill_category("technical", "Technical Skills")
                add_skill_category("soft", "Soft Skills")
                add_skill_category("languages", "Languages")
                add_skill_category("tools", "Tools & Technologies")
                add_block("SKILLS", skills_list)

            # Margins
            for section in doc.sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.7)
                section.right_margin = Inches(0.7)

            return doc

        except Exception:
            error = traceback.format_exc()
            save_error(error)